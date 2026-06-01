"""
create_template.py — создаёт report_template.docx с Jinja2-метками.
Без рисунков. Запускать при необходимости пересоздать шаблон.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Поля страницы ─────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(1.5)

# ── Стиль Normal ──────────────────────────────────────────────────────────────
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)


# ─── helpers ─────────────────────────────────────────────────────────────────
def p(text, bold=False, center=False, indent=True, size=12, italic=False):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after  = Pt(4)
    if indent and not center:
        par.paragraph_format.first_line_indent = Cm(1.25)
    r = par.add_run(text)
    r.bold    = bold
    r.italic  = italic
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    return par


def blank():
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after  = Pt(0)
    return par


def caption(text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after  = Pt(2)
    r = par.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    return par


def shd_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)


def cell(c, text, bold=False, center=False, bg=None, size=10):
    c.text = ''
    par = c.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    r = par.add_run(text)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    if bg:
        shd_cell(c, bg)


def set_col_widths(tbl, widths_cm):
    for row in tbl.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


# ─── Заголовок документа ─────────────────────────────────────────────────────
p('АКТ ТЕХНИЧЕСКОГО РАССЛЕДОВАНИЯ АВАРИИ', bold=True, center=True, indent=False, size=14)
blank()
p('Выемочный участок {{ t1_layer_name }}', bold=True, center=True, indent=False, size=13)
blank()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 1
# ════════════════════════════════════════════════════════════════════════════
p('1. ГОРНО-ГЕОЛОГИЧЕСКАЯ И ВЕНТИЛЯЦИОННАЯ ОБСТАНОВКА', bold=True, center=True, indent=False, size=12)
blank()

# ── Таблица 1 ─────────────────────────────────────────────────────────────────
p('{{ phrase_t1 }}', indent=True)
blank()
caption('Табл. 1  Основные параметры лавы {{ t1_layer_name }}')

t1 = doc.add_table(rows=1, cols=3)
t1.style     = 'Table Grid'
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
h = t1.rows[0].cells
cell(h[0], '№ п/п',                           bold=True, center=True, bg='1F4E79')
cell(h[1], 'Наименование параметра, ед. изм.', bold=True, center=True, bg='1F4E79')
cell(h[2], 'Величина',                         bold=True, center=True, bg='1F4E79')

rows_t1 = [
    ('1',  'Длина лавы, м',                                '{{ t1_length_m }}'),
    ('2',  'Длина отрабатываемого столба, м',              '{{ t1_pillar_length_m }}'),
    ('3',  'Нагрузка, т/сут',                              '{{ t1_daily_production }}'),
    ('4',  'Коэффициент извлечения угля, доли',            '{{ t1_coal_coef }}'),
    ('5',  'Отработка пласта',                             '{{ t1_plast_dev_type }}'),
    ('6',  'Способ управления кровлей',                    '{{ t1_conveyor_ctrl }}'),
    ('7',  'Породы кровли',                                '{{ t1_rock_type }}'),
    ('8',  'Схема проветривания',                          '{{ t1_vent_scheme }}'),
    ('9',  'Тип горно-шахтного комплекса',                 '{{ t1_complex_model }}'),
    ('10', 'Сечение вент. выработки, м²',                  '{{ t1_cross_sec_vent }}'),
    ('11', 'Сечение очистной выработки, м²',               '{{ t1_cross_sec_lava }}'),
    ('12', 'Допустимая скорость воздуха (min/max), м/с',   '{{ t1_air_speed_norm }}'),
    ('13', 'Скорость конвейера в очистной выработке, м/с', '{{ t1_conveyor_speed_lava }}'),
    ('14', 'Максимальное число людей на участке',           '{{ t1_employees_count }}'),
    ('15', 'Вынимаемая мощность пласта, м',                '{{ t1_thickness_m }}'),
    ('16', 'Плотность горной массы пласта, т/м³',          '{{ t1_rock_mass_density }}'),
    ('17', 'Плотность угля пласта, т/м³',                  '{{ t1_coal_density }}'),
]
for num, name, val in rows_t1:
    row = t1.add_row().cells
    cell(row[0], num,  center=True)
    cell(row[1], name)
    cell(row[2], val,  center=True)

set_col_widths(t1, [1.2, 10.3, 3.5])
blank()

# ── Таблица 2: Газовый баланс ─────────────────────────────────────────────────
p('{{ phrase_t2 }}', indent=True)
blank()
caption('Табл. 2  Расчётный газовый баланс выемочного участка')

t2 = doc.add_table(rows=1, cols=2)
t2.style     = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
h2 = t2.rows[0].cells
cell(h2[0], 'Наименование',              bold=True, center=True, bg='1F4E79')
cell(h2[1], 'Метановыделение, м³/т',     bold=True, center=True, bg='1F4E79')

# jinja2 loop rows
r_for = t2.add_row().cells
cell(r_for[0], '{%tr for r in t2_gas_rows %}')
cell(r_for[1], '')
r_dat = t2.add_row().cells
cell(r_dat[0], '{{ r.name }}')
cell(r_dat[1], '{{ r.value }}', center=True)
r_end = t2.add_row().cells
cell(r_end[0], '{%tr endfor %}')
cell(r_end[1], '')

set_col_widths(t2, [11.5, 3.5])
blank()

# ── Таблица 3: Вентиляция ─────────────────────────────────────────────────────
p('{{ phrase_t3 }}', indent=True)
p('{{ phrase_t3_vent }}', indent=True)
blank()
caption('Табл. 3  Параметры проветривания выемочного участка')

t3 = doc.add_table(rows=1, cols=3)
t3.style     = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
h3 = t3.rows[0].cells
cell(h3[0], '№',             bold=True, center=True, bg='1F4E79')
cell(h3[1], 'Наименование',  bold=True, center=True, bg='1F4E79')
cell(h3[2], 'Величина',      bold=True, center=True, bg='1F4E79')

rows_t3 = [
    ('1а', 'Расход воздуха на участке, м³/мин',                              '{{ t3_flow_uchastok }}'),
    ('1б', 'Расход воздуха в очистной выработке, м³/мин',                    '{{ t3_flow_lava }}'),
    ('2',  'Коэффициент утечек через выработанное пространство',              '{{ t3_leakage_coef }}'),
    ('3',  'Коэффициент учёта движения воздуха у призабойного пространства',  '{{ t3_distr_coef }}'),
    ('4а', 'Скорость воздуха на участке, м/с',                               '{{ t3_velocity_uchastok }}'),
    ('4б', 'Скорость воздуха в нижней части лавы, м/с',                      '{{ t3_velocity_lava }}'),
]
for num, name, val in rows_t3:
    row = t3.add_row().cells
    cell(row[0], num,  center=True)
    cell(row[1], name)
    cell(row[2], val,  center=True)

set_col_widths(t3, [1.2, 10.3, 3.5])
blank()

# ── Таблица 4: Атмосферное давление ──────────────────────────────────────────
p('{{ phrase_t4 }}', indent=True)
blank()
caption('Табл. 4  Атмосферное давление за {{ t4_date1 }} и {{ t4_date2 }}, мм рт. ст.')

t4 = doc.add_table(rows=2, cols=4)
t4.style     = 'Table Grid'
t4.alignment = WD_TABLE_ALIGNMENT.CENTER

# Строка-заголовок 1 (даты)
h4a = t4.rows[0].cells
cell(h4a[0], 'Время', bold=True, center=True, bg='1F4E79')
cell(h4a[1], '{{ t4_date1 }}', bold=True, center=True, bg='1F4E79')
cell(h4a[2], 'Время', bold=True, center=True, bg='1F4E79')
cell(h4a[3], '{{ t4_date2 }}', bold=True, center=True, bg='1F4E79')

# Строка-заголовок 2 (подпись «Давление»)
h4b = t4.rows[1].cells
cell(h4b[0], '',                       bold=True, center=True, bg='3A7DC9', size=9)
cell(h4b[1], 'Давление, мм рт. ст.',  bold=True, center=True, bg='3A7DC9', size=9)
cell(h4b[2], '',                       bold=True, center=True, bg='3A7DC9', size=9)
cell(h4b[3], 'Давление, мм рт. ст.',  bold=True, center=True, bg='3A7DC9', size=9)

r4_for = t4.add_row().cells
cell(r4_for[0], '{%tr for r in t4_pressure_rows %}', size=9)
for j in range(1, 4): cell(r4_for[j], '', size=9)

r4_dat = t4.add_row().cells
cell(r4_dat[0], '{{ r.time }}', center=True, size=9)
cell(r4_dat[1], '{{ r.p1 }}',  center=True, size=9)
cell(r4_dat[2], '{{ r.time }}', center=True, size=9)
cell(r4_dat[3], '{{ r.p2 }}',  center=True, size=9)

r4_end = t4.add_row().cells
cell(r4_end[0], '{%tr endfor %}', size=9)
for j in range(1, 4): cell(r4_end[j], '', size=9)

set_col_widths(t4, [2.5, 4.5, 2.5, 4.5])
blank()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 2
# ════════════════════════════════════════════════════════════════════════════
p('2. РЕЗУЛЬТАТЫ ИНСТРУМЕНТАЛЬНЫХ ИЗМЕРЕНИЙ', bold=True, center=True, indent=False, size=12)
blank()

# ── Таблица 1 раздела 2: CH4 на почве ────────────────────────────────────────
p('{{ phrase_ch4_floor }}', indent=True)
blank()
caption('Табл. 1  Результаты измерений концентрации метана на почве лавы')

t5 = doc.add_table(rows=1, cols=6)
t5.style     = 'Table Grid'
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
h5 = t5.rows[0].cells
for j, txt in enumerate(['№ секции', 'СН₄, %', '№ секции', 'СН₄, %', '№ секции', 'СН₄, %']):
    cell(h5[j], txt, bold=True, center=True, bg='1F4E79', size=9)

r5_for = t5.add_row().cells
cell(r5_for[0], '{%tr for r in ch4_floor_rows %}', size=9)
for j in range(1, 6): cell(r5_for[j], '', size=9)

r5_dat = t5.add_row().cells
for j, key in enumerate(['{{ r.sec1 }}', '{{ r.ch4_1 }}',
                          '{{ r.sec2 }}', '{{ r.ch4_2 }}',
                          '{{ r.sec3 }}', '{{ r.ch4_3 }}']):
    cell(r5_dat[j], key, center=True, size=9)

r5_end = t5.add_row().cells
cell(r5_end[0], '{%tr endfor %}', size=9)
for j in range(1, 6): cell(r5_end[j], '', size=9)

set_col_widths(t5, [2.5, 2.5, 2.5, 2.5, 2.5, 2.5])
blank()

# ── Таблица 2 раздела 2: CH4 среднее сечение ─────────────────────────────────
p('{{ phrase_ch4_mid }}', indent=True)
blank()
caption('Табл. 2  Концентрация метана в среднем сечении лавы')

t6 = doc.add_table(rows=1, cols=2)
t6.style     = 'Table Grid'
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
h6 = t6.rows[0].cells
cell(h6[0], '№ секции',              bold=True, center=True, bg='1F4E79')
cell(h6[1], 'Концентрация СН₄, %',  bold=True, center=True, bg='1F4E79')

r6_for = t6.add_row().cells
cell(r6_for[0], '{%tr for r in ch4_mid_rows %}')
cell(r6_for[1], '')
r6_dat = t6.add_row().cells
cell(r6_dat[0], '{{ r.section_no }}', center=True)
cell(r6_dat[1], '{{ r.ch4_percent }}', center=True)
r6_end = t6.add_row().cells
cell(r6_end[0], '{%tr endfor %}')
cell(r6_end[1], '')

set_col_widths(t6, [4.5, 4.5])
blank()

# ── Таблица 5: Сейсмические события ──────────────────────────────────────────
p('{{ phrase_seismic }}', indent=True)
blank()
caption('Табл. 5  Землетрясения за период {{ t4_date1 }}–{{ t4_date2 }}')

t7 = doc.add_table(rows=1, cols=8)
t7.style     = 'Table Grid'
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
h7 = t7.rows[0].cells
for j, txt in enumerate(['ID', 'Дата', 'Время (UTC)', 'Широта, °',
                          'Долгота, °', 'Глубина, км', 'Класс K', 'Магнитуда M']):
    cell(h7[j], txt, bold=True, center=True, bg='1F4E79', size=9)

r7_for = t7.add_row().cells
cell(r7_for[0], '{%tr for r in seismic_rows %}', size=9)
for j in range(1, 8): cell(r7_for[j], '', size=9)

r7_dat = t7.add_row().cells
for j, key in enumerate(['{{ r.event_id }}', '{{ r.date }}', '{{ r.time }}',
                          '{{ r.lat }}', '{{ r.lon }}', '{{ r.depth }}',
                          '{{ r.cls }}', '{{ r.mag }}']):
    cell(r7_dat[j], key, center=True, size=9)

r7_end = t7.add_row().cells
cell(r7_end[0], '{%tr endfor %}', size=9)
for j in range(1, 8): cell(r7_end[j], '', size=9)

set_col_widths(t7, [1.0, 2.4, 2.4, 1.7, 1.7, 1.8, 1.5, 2.0])

# ─────────────────────────────────────────────────────────────────────────────
doc.save('/home/claude/report_module/report_template.docx')
print('✓ report_template.docx создан (без рисунков)')
